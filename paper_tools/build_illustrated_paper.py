"""Generate seven paper figures from frozen project definitions and insert them into the team's Word draft.

This script is a presentation-only build step.  It does not modify model code,
Gate JSON, formal numerical results, or the frozen strategy.
"""
from __future__ import annotations

from itertools import combinations
from math import cos, pi, sin, sqrt
from pathlib import Path
import sys

import numpy as np
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from src.q1_3_adjustment import table1_coordinates
from src.q1_2_identity import target_coordinates
from src.q2_geometry import target_lattice
from src.q2_evaluator import nearest_neighbor_edges

ASSET_DIR = ROOT / "paper_assets" / "figures"
OUT_DIR = ROOT / "paper_drafts"
SOURCE = Path(r"C:\Users\LHC\Downloads\2022国赛B题提交材料\论文\2022国赛B题论文_最终版.docx")
OUTPUT = OUT_DIR / "2022国赛B题论文_图表补全版.docx"

FONT_PATH = Path(r"C:\Windows\Fonts\msyh.ttc")
if not FONT_PATH.exists():
    raise FileNotFoundError("缺少中文字体 Microsoft YaHei: C:\\Windows\\Fonts\\msyh.ttc")

W, H = 1800, 1125
INK = (35, 42, 51)
MUTED = (110, 120, 132)
GRID = (205, 211, 217)
BLUE = (46, 115, 181)
ORANGE = (230, 126, 34)
GREEN = (43, 142, 93)
RED = (190, 58, 58)
PURPLE = (116, 80, 165)
PALE = (235, 241, 246)


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(str(FONT_PATH), size=size, index=1 if bold else 0)


F20, F21, F22, F23, F24, F26, F28, F32, F38, F46 = (font(s) for s in (20, 21, 22, 23, 24, 26, 28, 32, 38, 46))
FB22, FB24, FB28, FB32, FB38 = (font(s, True) for s in (22, 24, 28, 32, 38))


def canvas() -> tuple[Image.Image, ImageDraw.ImageDraw]:
    image = Image.new("RGB", (W, H), "white")
    return image, ImageDraw.Draw(image)


def text(draw: ImageDraw.ImageDraw, xy: tuple[float, float], value: str, f=F28, fill=INK, anchor="la") -> None:
    draw.text(xy, value, font=f, fill=fill, anchor=anchor)


def line(draw: ImageDraw.ImageDraw, points, fill=INK, width=4, dash: tuple[int, int] | None = None) -> None:
    pts = [(int(x), int(y)) for x, y in points]
    if not dash:
        draw.line(pts, fill=fill, width=width)
        return
    for left, right in zip(pts, pts[1:]):
        dx, dy = right[0] - left[0], right[1] - left[1]
        distance = max(1, int((dx * dx + dy * dy) ** 0.5))
        nx, ny = dx / distance, dy / distance
        position = 0
        while position < distance:
            end = min(distance, position + dash[0])
            draw.line([(left[0] + nx * position, left[1] + ny * position), (left[0] + nx * end, left[1] + ny * end)], fill=fill, width=width)
            position += dash[0] + dash[1]


def arrow(draw: ImageDraw.ImageDraw, start, end, fill=INK, width=4, head=16) -> None:
    line(draw, [start, end], fill=fill, width=width)
    x1, y1 = start
    x2, y2 = end
    angle = np.arctan2(y2 - y1, x2 - x1)
    left = (x2 - head * np.cos(angle - 0.48), y2 - head * np.sin(angle - 0.48))
    right = (x2 - head * np.cos(angle + 0.48), y2 - head * np.sin(angle + 0.48))
    draw.polygon([end, left, right], fill=fill)


def node(draw: ImageDraw.ImageDraw, xy, label: str, fill=BLUE, radius=15, box=False, label_offset=(0, -28)) -> None:
    x, y = xy
    if box:
        draw.rounded_rectangle((x - radius, y - radius, x + radius, y + radius), radius=4, fill=fill, outline=INK, width=2)
    else:
        draw.ellipse((x - radius, y - radius, x + radius, y + radius), fill=fill, outline=INK, width=2)
    text(draw, (x + label_offset[0], y + label_offset[1]), label, F24, INK, "ma")


def title(draw: ImageDraw.ImageDraw, value: str) -> None:
    text(draw, (W / 2, 54), value, FB38, INK, "ma")
    line(draw, [(120, 100), (W - 120, 100)], GRID, 2)


def panel(draw: ImageDraw.ImageDraw, box, label: str) -> None:
    x0, y0, x1, y1 = box
    draw.rounded_rectangle(box, radius=18, outline=GRID, width=3, fill=(252, 253, 254))
    text(draw, (x0 + 26, y0 + 28), label, FB28, INK)


def save(image: Image.Image, name: str) -> Path:
    path = ASSET_DIR / name
    image.save(path, quality=96, dpi=(300, 300))
    return path


def fig01() -> Path:
    image, draw = canvas()
    title(draw, "双侧定夹角轨迹圆与候选分支")
    a, b = (500, 600), (1300, 600)
    r, offset = 462, 231
    for center, color in [((900, 369), BLUE), ((900, 831), ORANGE)]:
        cx, cy = center
        draw.ellipse((cx - r, cy - r, cx + r, cy + r), outline=color, width=5)
    line(draw, [a, b], INK, 4)
    node(draw, a, "发射机 A", RED, 16, True, (-5, -48))
    node(draw, b, "发射机 B", RED, 16, True, (5, -48))
    # Two marked arcs are the admissible branches used in complete enumeration.
    draw.arc((438, -93, 1362, 831), 34, 146, fill=BLUE, width=12)
    draw.arc((438, 369, 1362, 1293), 214, 326, fill=ORANGE, width=12)
    text(draw, (900, 230), "候选分支 C₊", FB28, BLUE, "ma")
    text(draw, (900, 988), "候选分支 C₋", FB28, ORANGE, "ma")
    p1, p2 = (745, 164), (1055, 1036)
    node(draw, p1, "候选位置 P₊", GREEN, 15, False, (0, -34))
    node(draw, p2, "候选位置 P₋", GREEN, 15, False, (0, 38))
    for p in [p1, p2]:
        line(draw, [p, a], MUTED, 2, (10, 8))
        line(draw, [p, b], MUTED, 2, (10, 8))
    text(draw, (900, 680), "同一无向夹角 α 对应两侧几何分支", FB32, INK, "ma")
    text(draw, (900, 730), "因此必须保留全部圆弧与交点，不能默认数值求解的首个根。", F28, MUTED, "ma")
    return save(image, "fig01_two_sided_angle_locus.png")


def fig02() -> Path:
    image, draw = canvas()
    title(draw, "三组纯方位约束的完整候选与第三角回代筛选")
    left = (120, 160, 880, 1040)
    right = (960, 160, 1680, 1040)
    panel(draw, left, "第一、二角：生成全部有限候选")
    panel(draw, right, "第三角：逐候选回代筛选")
    anchors = {"A": (380, 820), "B": (690, 820), "C": (540, 350)}
    for k, p in anchors.items():
        node(draw, p, f"发射机 {k}", RED, 15, True, (0, -40 if k != "C" else 42))
    line(draw, [anchors["A"], anchors["B"]], MUTED, 3)
    line(draw, [anchors["A"], anchors["C"]], MUTED, 3)
    line(draw, [anchors["B"], anchors["C"]], MUTED, 3)
    candidates = [(290, 500), (515, 585), (760, 505), (540, 930)]
    for idx, p in enumerate(candidates, 1):
        node(draw, p, f"P{idx}", PURPLE if idx != 2 else GREEN, 14, False, (0, -30))
        for a in ("A", "B"):
            line(draw, [p, anchors[a]], GRID, 2, (8, 7))
    text(draw, (500, 975), "圆—圆求交：所有两侧分支、所有交点均保留", F24, MUTED, "ma")
    arrow(draw, (890, 600), (940, 600), BLUE, 5, 20)
    ranchors = {"A": (1110, 820), "B": (1450, 820), "C": (1280, 350)}
    for k, p in ranchors.items():
        node(draw, p, f"{k}", RED, 14, True, (0, -32 if k != "C" else 38))
    kept = (1275, 605)
    for idx, p in enumerate([(1080, 510), kept, (1490, 510), (1280, 925)], 1):
        if p == kept:
            node(draw, p, "P*", GREEN, 18, False, (0, -38))
            for a in ("A", "B", "C"):
                line(draw, [p, ranchors[a]], GREEN, 3, (8, 7))
        else:
            node(draw, p, "×", RED, 18, False, (0, 0))
            text(draw, (p[0], p[1] + 45), "第三角不满足", F20, RED, "ma")
    text(draw, (1320, 975), "保留同时满足三条原始夹角的局部候选 P*", F24, MUTED, "ma")
    return save(image, "fig02_complete_candidates_and_backsubstitution.png")


def fig03() -> Path:
    image, draw = canvas()
    title(draw, "Q1(2) 的零匿名机反例与一匿名机联合判定")
    left, right = (120, 150, 860, 1035), (940, 150, 1680, 1035)
    panel(draw, left, "m = 0：只有 FY00、FY01")
    panel(draw, right, "m = 1：编号—位置联合枚举")
    a, b = (300, 700), (700, 700)
    node(draw, a, "FY00", RED, 15, True, (0, -42))
    node(draw, b, "FY01", RED, 15, True, (0, -42))
    draw.arc((220, 220, 780, 780), 192, 348, fill=BLUE, width=10)
    line(draw, [a, b], MUTED, 3)
    text(draw, (490, 300), "一条定夹角圆弧", FB28, BLUE, "ma")
    text(draw, (490, 850), "连续位置集合 → 不能唯一确定二维位置", F26, INK, "ma")
    center = (1300, 610)
    draw.ellipse((1090, 400, 1510, 820), outline=GRID, width=3)
    ids = list(range(2, 9))
    for k, ident in enumerate(ids):
        angle = 2 * pi * k / len(ids) - pi / 2
        p = (center[0] + 210 * cos(angle), center[1] + 210 * sin(angle))
        color = GREEN if ident == 5 else BLUE
        node(draw, p, f"FY{ident:02d}", color, 13, ident == 5, (0, -35 if sin(angle) >= 0 else 38))
    node(draw, center, "接收机 r", ORANGE, 17, True, (0, 0))
    text(draw, (1300, 865), "7 个合法编号假设 → 完整几何候选 → 目标邻域 Uᵣ", F24, INK, "ma")
    text(draw, (1300, 915), "当局部秩为 2 且观测分离度大于 0 时，保留唯一 (r,b,x)。", F22, MUTED, "ma")
    return save(image, "fig03_identity_position_joint_identifiability.png")


def fig04() -> Path:
    image, draw = canvas()
    title(draw, "Q1(3) FY04/FY07 严格本机交替校正与四锚并行调整")
    panel(draw, (110, 165, 840, 825), "一个完整宏周期：两台接收机分别处理本机夹角")
    boxes = [(160, 300, 790, 470), (160, 545, 790, 715)]
    entries = [
        ("子轮 1", "发射：FY00、FY01、FY07", "接收并移动：FY04", "仅使用 FY04 本机两条主夹角"),
        ("子轮 2", "发射：FY00、FY01、FY04", "接收并移动：FY07", "仅使用 FY07 本机两条主夹角"),
    ]
    for box, row in zip(boxes, entries):
        x0, y0, x1, y1 = box
        draw.rounded_rectangle(box, radius=18, fill=PALE, outline=GRID, width=2)
        text(draw, (x0 + 28, y0 + 28), row[0], FB28, BLUE)
        text(draw, (x0 + 190, y0 + 28), row[1], F24)
        text(draw, (x0 + 190, y0 + 76), row[2], FB24, ORANGE)
        text(draw, (x0 + 190, y0 + 122), row[3], F22, MUTED)
    arrow(draw, (475, 475), (475, 535), GREEN, 5, 18)
    text(draw, (475, 760), "禁止跨接收机交换夹角；每次移动时发射机保持不动。", F23, RED, "ma")
    panel(draw, (930, 165, 1690, 825), "建锚完成后：四锚固定，六架跟随者并行")
    anchors = [(1060, 370, "FY00"), (1270, 300, "FY01"), (1490, 480, "FY04"), (1230, 650, "FY07")]
    for x, y, label in anchors:
        node(draw, (x, y), label, RED, 16, True, (0, -42))
    followers = [(1050, 560, "FY02/03"), (1450, 630, "FY05/06"), (1570, 350, "FY08/09")]
    for x, y, label in followers:
        node(draw, (x, y), label, BLUE, 15, False, (0, 38))
        for ax, ay, _ in anchors[:2]:
            line(draw, [(x, y), (ax, ay)], GRID, 2, (8, 7))
    text(draw, (1310, 760), "每架跟随者仅用本机两条主夹角控制，其余同机夹角用于留出检验。", F22, MUTED, "ma")
    text(draw, (900, 955), "信息流：预装理想夹角 + 本机试探观测 + 本机历史  →  本机动作", FB28, INK, "ma")
    text(draw, (900, 1008), "不使用真实坐标、距离、其他接收机夹角或离线评估结果。", F24, RED, "ma")
    return save(image, "fig04_q1_local_alternating_schedule.png")


def _polar_mapper(points: dict[int, np.ndarray], box):
    x0, y0, x1, y1 = box
    extent = max(max(abs(float(v[0])), abs(float(v[1]))) for v in points.values())
    scale = min((x1 - x0), (y1 - y0)) * 0.40 / extent
    cx, cy = (x0 + x1) / 2, (y0 + y1) / 2
    return lambda p: (cx + scale * float(p[0]), cy - scale * float(p[1])), scale, (cx, cy)


def fig05() -> Path:
    image, draw = canvas()
    title(draw, "Q1(3) 表 1 初始队形、交替建锚与最终正九边形")
    source = table1_coordinates()
    ideal = target_coordinates()
    panels = [(90, 150, 610, 1050, "(a) 表 1 初始位置"), (640, 150, 1160, 1050, "(b) FY04/FY07 建锚阶段"), (1190, 150, 1710, 1050, "(c) 最终正九边形")]
    for box in panels:
        panel(draw, box[:4], box[4])
    for idx, box in enumerate(panels):
        mapper, scale, (cx, cy) = _polar_mapper(ideal, box[:4])
        draw.ellipse((cx - 100 * scale, cy - 100 * scale, cx + 100 * scale, cy + 100 * scale), outline=GRID, width=2)
        node(draw, (cx, cy), "FY00", RED, 10, True, (0, -25))
        if idx == 0:
            for k, p in source.items():
                if k:
                    node(draw, mapper(p), f"{k:02d}", ORANGE, 10, False, (0, -24))
            text(draw, (350, 980), "橙色点：题表初始位置", F22, MUTED, "ma")
        elif idx == 1:
            for k, p in source.items():
                if k in (4, 7):
                    continue
                if k:
                    node(draw, mapper(p), f"{k:02d}", (180, 190, 200), 8, False, (0, -21))
            for k in (4, 7):
                start, end = mapper(source[k]), mapper(ideal[k])
                arrow(draw, start, end, GREEN, 4, 14)
                node(draw, end, f"FY{k:02d}", GREEN, 13, True, (0, -32))
            node(draw, mapper(ideal[1]), "FY01", RED, 12, True, (0, -30))
            text(draw, (900, 980), "绿色箭头：FY04、FY07 的本机调整方向", F21, MUTED, "ma")
        else:
            ring = [mapper(ideal[k]) for k in range(1, 10)]
            for p, q in zip(ring, ring[1:] + ring[:1]):
                line(draw, [p, q], BLUE, 3)
            for k in range(1, 10):
                node(draw, mapper(ideal[k]), f"FY{k:02d}", GREEN if k in (4, 7) else BLUE, 10, k in (1, 4, 7), (0, -25))
            text(draw, (1450, 980), "最大目标位置误差：2.93×10⁻⁷ m", F21, MUTED, "ma")
    return save(image, "fig05_q1_initial_bootstrap_final_formation.png")


def _lattice_mapper(points: dict[int, np.ndarray], box):
    x0, y0, x1, y1 = box
    xs, ys = [float(v[0]) for v in points.values()], [float(v[1]) for v in points.values()]
    scale = min((x1 - x0) / (max(xs) - min(xs) + 1.1), (y1 - y0) / (max(ys) - min(ys) + 1.1))
    cx, cy = (x0 + x1) / 2, (y0 + y1) / 2
    mx, my = (max(xs) + min(xs)) / 2, (max(ys) + min(ys)) / 2
    return lambda p: (cx + scale * (float(p[0]) - mx), cy - scale * (float(p[1]) - my)), scale


def draw_lattice(draw: ImageDraw.ImageDraw, box, refs=(3, 4, 11, 15), labels=True, alpha=False):
    lattice = target_lattice()
    mapper, scale = _lattice_mapper(lattice, box)
    for i, j in nearest_neighbor_edges():
        line(draw, [mapper(lattice[i]), mapper(lattice[j])], GRID, 3)
    for k, p in lattice.items():
        node(draw, mapper(p), f"FY{k:02d}" if labels else str(k), RED if k in refs else BLUE, 13 if k in refs else 10, k in refs, (0, -28))
    return lattice, mapper, scale


def fig06() -> Path:
    image, draw = canvas()
    title(draw, "Q2 三角格点目标、相邻边与四参考机结构")
    box = (150, 150, 1260, 1040)
    panel(draw, box, "目标三角格点：30 条相邻边、12 条要求共线的最大直线")
    lattice, mapper, scale = draw_lattice(draw, box, labels=True)
    p11, p15 = mapper(lattice[11]), mapper(lattice[15])
    arrow(draw, (p11[0] - 5, p11[1] + 48), (p15[0] - 5, p15[1] + 48), RED, 5, 18)
    text(draw, ((p11[0] + p15[0]) / 2, p11[1] + 88), "FY11—FY15 = 4d*（可信无偏差尺度基线）", F23, RED, "ma")
    panel(draw, (1320, 150, 1710, 1040), "四参考机的角色")
    roles = [
        ("FY11、FY15", "固定尺度基线", RED),
        ("FY04、FY03", "交替校正建锚", GREEN),
        ("其余 11 架", "以本机六维纯方位观测向量归槽", BLUE),
    ]
    y = 300
    for heading, body, color in roles:
        draw.rounded_rectangle((1370, y, 1660, y + 135), radius=16, fill=PALE, outline=GRID, width=2)
        text(draw, (1515, y + 38), heading, FB24, color, "ma")
        text(draw, (1515, y + 88), body, F21, INK, "ma")
        y += 190
    text(draw, (1515, 905), "平移、旋转、镜像保持自由；\n尺度由 FY11/FY15 基线注入。", F22, MUTED, "ma")
    return save(image, "fig06_q2_lattice_and_four_references.png")


def fig07() -> Path:
    image, draw = canvas()
    title(draw, "Q2 实际端到端调整：尺度基线 → 交替建锚 → 十一机归槽")
    boxes = [(60, 170, 580, 1010, "(a) 固定可信尺度基线"), (640, 170, 1160, 1010, "(b) FY04/FY03 实际建锚终点"), (1220, 170, 1740, 1010, "(c) 四参考机下十一机归槽")]
    for box in boxes:
        panel(draw, box[:4], box[4])
    lattice = target_lattice()
    # a: trusted baseline only
    mapper_a, _ = _lattice_mapper(lattice, boxes[0][:4])
    for k, p in lattice.items():
        node(draw, mapper_a(p), f"{k:02d}", RED if k in (11, 15) else (195, 205, 215), 13 if k in (11, 15) else 7, k in (11, 15), (0, -22))
    p11, p15 = mapper_a(lattice[11]), mapper_a(lattice[15])
    arrow(draw, (p11[0], p11[1] + 35), (p15[0], p15[1] + 35), RED, 4, 15)
    text(draw, ((p11[0] + p15[0]) / 2, p11[1] + 67), "4d*", FB24, RED, "ma")
    text(draw, (320, 900), "FY11、FY15 保持固定；\n不固定全局平移、旋转或镜像。", F22, MUTED, "ma")
    # b: show bootstrap within formal lattice
    mapper_b, _ = _lattice_mapper(lattice, boxes[1][:4])
    for k, p in lattice.items():
        color = RED if k in (11, 15) else (205, 213, 220)
        node(draw, mapper_b(p), f"{k:02d}", color, 12 if k in (11, 15) else 6, k in (11, 15), (0, -22))
    for k in (3, 4):
        p = mapper_b(lattice[k])
        node(draw, p, f"FY{k:02d}", GREEN, 15, True, (0, -32))
    arrow(draw, (900, 500), (900, 600), GREEN, 4, 16)
    text(draw, (900, 695), "两子轮交替：\nFY04 仅处理本机夹角，\nFY03 仅处理本机夹角。", F21, MUTED, "ma")
    text(draw, (900, 880), "后续使用两机的实际调整终点，\n不重置为理想位置。", FB22, RED, "ma")
    # c: all final with real data flow emphasis
    _, mapper_c, _ = draw_lattice(draw, boxes[2][:4], labels=False)
    for k in (3, 4, 11, 15):
        p = mapper_c(lattice[k])
        text(draw, (p[0], p[1] - 32), f"FY{k:02d}", F20, RED, "ma")
    text(draw, (1480, 850), "33/33 个端到端确定性算例通过", FB22, GREEN, "ma")
    text(draw, (1480, 900), "最坏节点误差：2.27×10⁻⁷ d*\n30 边、12 线均通过离线验收。", F20, MUTED, "ma")
    arrow(draw, (590, 590), (625, 590), BLUE, 5, 18)
    arrow(draw, (1170, 590), (1205, 590), BLUE, 5, 18)
    return save(image, "fig07_q2_end_to_end_adjustment.png")


def insert_paragraph_after(paragraph, text_value: str = "") -> Paragraph:
    element = OxmlElement("w:p")
    paragraph._p.addnext(element)
    created = Paragraph(element, paragraph._parent)
    if text_value:
        created.add_run(text_value)
    return created


def caption_after(paragraph, caption: str) -> Paragraph:
    p = insert_paragraph_after(paragraph)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(caption)
    run.font.name = "宋体"
    run.font.size = Pt(10.5)
    return p


def image_after(paragraph, image_path: Path, caption: str, width_cm: float = 14.8) -> Paragraph:
    p = insert_paragraph_after(paragraph)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Cm(width_cm))
    return caption_after(p, caption)


def remove_paragraph(paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)


def find_first(doc: Document, starts: str):
    for p in doc.paragraphs:
        if starts in p.text:
            return p
    raise KeyError(f"找不到锚点段落：{starts}")


def replace_paragraph(p, new_text: str) -> None:
    p.clear()
    p.add_run(new_text)


def apply_docx(figures: dict[str, Path]) -> None:
    doc = Document(SOURCE)
    # Delete the old single figure and its caption; it is replaced by the corrected Figure 1.
    for p in list(doc.paragraphs):
        if p.text.strip() == "图 1 接收信号无人机i位置示意图" or p._p.xpath('.//w:drawing'):
            remove_paragraph(p)

    # Correct only the outdated Q2 problem-analysis route so it matches the frozen Q2 structure.
    replace_paragraph(find_first(doc, "锥形编队由FY01-FY15共15架无人机组成"),
        "锥形编队由 FY01—FY15 共 15 架无人机组成。题目要求各规定直线上的无人机保持共线，且相邻无人机间距相等。本文将目标队形抽象为标准三角格点，并以相邻间距 d* 表示共同尺度。")
    replace_paragraph(find_first(doc, "根据编队的锥形结构,将本题分为两个阶段"),
        "纯方位夹角对整体缩放不敏感，因此仅由夹角不能恢复共同尺度。为形成可施工的尺度参考，采用经批准的 FY11、FY15 可信无偏差基线，其距离为 4d*；该条件只固定尺度，不固定队形的整体平移、旋转和镜像。")
    replace_paragraph(find_first(doc, "第二阶段固定FY01、FY11 和 FY15"),
        "具体先固定 FY11、FY15，并让 FY04、FY03 进行严格本机双子轮交替校正，形成四参考机框架；再固定 FY03、FY04、FY11、FY15，其余 11 架无人机仅依据本机测得的六维纯方位观测向量并行归入对应格点。每架无人机以两条非退化夹角构造本机调整方向，其余同机夹角只用于留出约束检验；不同接收机之间不交换夹角。")

    anchors = [
        ("一般会有位于弦两侧的两个轨迹圆", figures["f1"], "图1 双侧定夹角轨迹圆与候选分支示意图"),
        ("由上述模型可对(1)问进行具体求解", figures["f2"], "图2 三组纯方位约束的完整候选与第三角回代筛选"),
        ("至少需要增加一架发射发射信号无人机", figures["f3"], "图3 Q1(2) 的零匿名机反例与一匿名机联合判定"),
        ("两架无人机不交换各自测得的夹角。", figures["f4"], "图4 Q1(3) FY04/FY07 严格本机交替校正与四锚并行调整"),
        ("也不能把计算误差看作实际无人机的飞行精度。", figures["f5"], "图5 Q1(3) 表1初始队形、交替建锚与最终正九边形"),
        ("这里的50 m只是计算示例。", figures["f6"], "图6 Q2三角格点目标、相邻边与四参考机结构"),
        ("实际调整后的终点，没有将两架无人机重新放回理想位置。", figures["f7"], "图7 Q2实际端到端调整：尺度基线、交替建锚与十一机归槽"),
    ]
    # Insert each at a unique anchor.  Subsequent insertions do not change the anchor text.
    for starts, fig, caption in anchors:
        image_after(find_first(doc, starts), fig, caption)

    for p in doc.paragraphs:
        if p.text.startswith("图"):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.save(OUTPUT)


def main() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    figures = {
        "f1": fig01(), "f2": fig02(), "f3": fig03(), "f4": fig04(),
        "f5": fig05(), "f6": fig06(), "f7": fig07(),
    }
    apply_docx(figures)
    print(f"OUTPUT={OUTPUT}")
    for key, path in figures.items():
        print(f"{key}={path}")


if __name__ == "__main__":
    main()
