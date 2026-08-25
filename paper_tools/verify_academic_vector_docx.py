"""Mechanical integrity checks for the academic-vector DOCX artifact."""

from __future__ import annotations

import argparse
import hashlib
from io import BytesIO
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from lxml import etree
from PIL import Image


NS = {
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "asvg": "http://schemas.microsoft.com/office/drawing/2016/SVG/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "pr": "http://schemas.openxmlformats.org/package/2006/relationships",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
}


def verify(docx_path: Path, fallback_directory: Path) -> None:
    with ZipFile(docx_path) as archive:
        if archive.testzip() is not None:
            raise AssertionError("DOCX ZIP integrity check failed")
        names = set(archive.namelist())
        svg_media = sorted(
            name
            for name in names
            if name.startswith("word/media/") and name.endswith(".svg")
        )
        if len(svg_media) != 7:
            raise AssertionError(f"Expected 7 SVG media files, found {len(svg_media)}")

        document_xml = etree.fromstring(archive.read("word/document.xml"))
        relationships_xml = etree.fromstring(
            archive.read("word/_rels/document.xml.rels")
        )
        svg_blips = document_xml.xpath(".//asvg:svgBlip", namespaces=NS)
        if len(svg_blips) != 7:
            raise AssertionError(f"Expected 7 SVG references, found {len(svg_blips)}")

        relationships = {
            relationship.get("Id"): relationship.get("Target")
            for relationship in relationships_xml.xpath("./pr:Relationship", namespaces=NS)
        }
        for svg_blip in svg_blips:
            relationship_id = svg_blip.get(f"{{{NS['r']}}}embed")
            target = relationships.get(relationship_id)
            if not target or f"word/{target}" not in names:
                raise AssertionError(f"Broken SVG relationship: {relationship_id}")

        extents = document_xml.xpath(".//wp:inline/wp:extent", namespaces=NS)
        if len(extents) != 7:
            raise AssertionError(f"Expected 7 inline extents, found {len(extents)}")
        expected_ratio = 1200 / 750
        for index, extent in enumerate(extents, start=1):
            ratio = int(extent.get("cx")) / int(extent.get("cy"))
            if abs(ratio - expected_ratio) > 0.01:
                raise AssertionError(
                    f"Figure {index} is distorted: extent ratio {ratio:.4f}"
                )

        fallback_paths = sorted(fallback_directory.glob("fig*.png"))
        if len(fallback_paths) != 7:
            raise AssertionError("Expected 7 fallback PNG assets")
        for index, fallback_path in enumerate(fallback_paths, start=1):
            package_name = f"word/media/image{index}.png"
            package_bytes = archive.read(package_name)
            source_bytes = fallback_path.read_bytes()
            if hashlib.sha256(package_bytes).digest() != hashlib.sha256(source_bytes).digest():
                raise AssertionError(f"Fallback PNG mismatch for figure {index}")
            with Image.open(BytesIO(package_bytes)) as image:
                if image.size != (1200, 750):
                    raise AssertionError(
                        f"Unexpected PNG size for figure {index}: {image.size}"
                    )

    document = Document(docx_path)
    if len(document.inline_shapes) != 7:
        raise AssertionError("python-docx did not recover all seven inline figures")

    drawing_paragraphs: list[int] = []
    for paragraph_index, paragraph in enumerate(document.paragraphs):
        if paragraph._p.xpath(".//w:drawing"):
            drawing_paragraphs.append(paragraph_index)
    if len(drawing_paragraphs) != 7:
        raise AssertionError("Expected exactly seven drawing paragraphs")

    captions: list[str] = []
    for figure_number, paragraph_index in enumerate(drawing_paragraphs, start=1):
        next_text = ""
        for following in document.paragraphs[paragraph_index + 1 :]:
            if following.text.strip():
                next_text = following.text.strip()
                break
        if not next_text.startswith(f"图{figure_number}"):
            raise AssertionError(
                f"Figure {figure_number} is not followed by its caption: {next_text!r}"
            )
        captions.append(next_text)

    print("DOCX_ZIP_INTEGRITY=PASS")
    print("SVG_VECTOR_MEDIA=7")
    print("PNG_COMPATIBILITY_FALLBACK=7")
    print("INLINE_FIGURES=7")
    print("ASPECT_RATIO_CHECK=PASS")
    print("CAPTION_ADJACENCY=PASS")
    for caption in captions:
        print(f"CAPTION={caption}")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("docx_path", type=Path)
    parser.add_argument("--fallback-dir", type=Path, required=True)
    arguments = parser.parse_args()
    verify(arguments.docx_path.resolve(), arguments.fallback_dir.resolve())


if __name__ == "__main__":
    main()
